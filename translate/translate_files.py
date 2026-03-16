#!/usr/bin/env python3
from __future__ import annotations

import argparse
import concurrent.futures
import importlib
import os
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
_OPENAI_THREAD_LOCAL = threading.local()


def eprint(message: str) -> None:
	print(message, file=sys.stderr)


def iter_files(input_dir: Path) -> Iterable[Path]:
	for path in input_dir.rglob("*"):
		if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
			yield path


def chunk_text(text: str, max_chars: int) -> list[str]:
	# Chunk conservatively by paragraphs to preserve structure and avoid huge prompts.
	paragraphs = [p.strip() for p in text.splitlines()]
	chunks: list[str] = []
	current: list[str] = []
	current_len = 0

	for p in paragraphs:
		if not p:
			# Preserve empty lines as separators.
			p = ""
		add_len = len(p) + 1
		if current and current_len + add_len > max_chars:
			chunks.append("\n".join(current).strip("\n"))
			current = []
			current_len = 0
		current.append(p)
		current_len += add_len

	if current:
		chunks.append("\n".join(current).strip("\n"))

	# Avoid empty chunks.
	return [c for c in chunks if c.strip()]


class Translator:
	def translate(self, *, text: str, source_lang: str, target_lang: str) -> str:
		raise NotImplementedError


@dataclass(frozen=True)
class OpenAITranslator(Translator):
	model: str

	def translate(self, *, text: str, source_lang: str, target_lang: str) -> str:
		api_key = os.getenv("OPENAI_API_KEY")
		if not api_key:
			raise RuntimeError("OPENAI_API_KEY is not set.")

		try:
			from openai import OpenAI
			from openai import APIError, APITimeoutError, RateLimitError
		except Exception as exc:  # pragma: no cover
			raise RuntimeError(
				"OpenAI SDK not installed. Install translate/requirements.txt."
			) from exc

		client = getattr(_OPENAI_THREAD_LOCAL, "client", None)
		if client is None:
			client = OpenAI(api_key=api_key)
			_OPENAI_THREAD_LOCAL.client = client

		system = (
			"You are a professional translator. Translate the provided text "
			f"from {source_lang} to {target_lang}. Preserve meaning and paragraph breaks. "
			"Do not add notes, disclaimers, or commentary. Output only the translated text."
		)

		delay_seconds = 0.5
		for attempt in range(6):
			try:
				resp = client.chat.completions.create(
					model=self.model,
					messages=[
						{"role": "system", "content": system},
						{"role": "user", "content": text},
					],
					temperature=0,
				)
				content = resp.choices[0].message.content
				if not content:
					return ""
				return content.strip()
			except (RateLimitError, APITimeoutError, APIError) as exc:
				if attempt == 5:
					raise
				# Basic exponential backoff to smooth out bursts when parallelizing.
				time.sleep(delay_seconds)
				delay_seconds = min(8.0, delay_seconds * 2)
			except Exception:
				raise

		return ""


@dataclass(frozen=True)
class ArgosTranslator(Translator):
	# Requires local Argos installation + sl->en model installed.
	def translate(self, *, text: str, source_lang: str, target_lang: str) -> str:
		try:
			argos_translate = importlib.import_module("argostranslate.translate")
		except Exception as exc:  # pragma: no cover
			raise RuntimeError(
				"Argos Translate not installed. Install it and an sl->en model."
			) from exc

		installed_languages = argos_translate.get_installed_languages()
		source = next((l for l in installed_languages if l.code == source_lang), None)
		if source is None:
			raise RuntimeError(
				f"Argos source language '{source_lang}' not installed (model missing)."
			)

		target = next((l for l in installed_languages if l.code == target_lang), None)
		if target is None:
			raise RuntimeError(
				f"Argos target language '{target_lang}' not installed (model missing)."
			)

		translation = source.get_translation(target)
		return translation.translate(text)


def get_translator(provider: str, model: str) -> Translator:
	if provider == "openai":
		return OpenAITranslator(model=model)
	if provider == "argos":
		return ArgosTranslator()
	raise ValueError(f"Unknown provider: {provider}")


def extract_text_from_pdf(path: Path) -> str:
	try:
		import pdfplumber
	except Exception as exc:  # pragma: no cover
		raise RuntimeError("pdfplumber not installed. Install translate/requirements.txt.") from exc

	parts: list[str] = []
	with pdfplumber.open(str(path)) as pdf:
		for page in pdf.pages:
			page_text = page.extract_text() or ""
			page_text = page_text.strip()
			if page_text:
				parts.append(page_text)
	return "\n\n".join(parts).strip()


def write_pdf_from_text(path: Path, text: str) -> None:
	# Simple “text PDF” output (not layout-preserving). Good for readable delivery.
	try:
		pagesizes = importlib.import_module("reportlab.lib.pagesizes")
		units = importlib.import_module("reportlab.lib.units")
		pdfmetrics = importlib.import_module("reportlab.pdfbase.pdfmetrics")
		ttfonts = importlib.import_module("reportlab.pdfbase.ttfonts")
		canvas_mod = importlib.import_module("reportlab.pdfgen.canvas")
	except Exception as exc:  # pragma: no cover
		raise RuntimeError("reportlab not installed. Install translate/requirements.txt.") from exc

	LETTER = pagesizes.LETTER
	inch = units.inch
	TTFont = ttfonts.TTFont
	Canvas = canvas_mod.Canvas

	# Try to use a Unicode-capable font if available (falls back to built-in if not).
	# On macOS, Arial Unicode may exist; if not, Helvetica will work for most ASCII.
	font_name = "Helvetica"
	for candidate_name, candidate_path in [
		("DejaVuSans", "/Library/Fonts/DejaVuSans.ttf"),
		("ArialUnicodeMS", "/Library/Fonts/Arial Unicode.ttf"),
	]:
		try:
			if Path(candidate_path).exists():
				pdfmetrics.registerFont(TTFont(candidate_name, candidate_path))
				font_name = candidate_name
				break
		except Exception:
			# If font registration fails, keep going with default.
			pass

	page_width, page_height = LETTER
	margin = 0.75 * inch
	max_width = page_width - (2 * margin)
	line_height = 12
	font_size = 10

	c = Canvas(str(path), pagesize=LETTER)
	c.setTitle(path.stem)
	c.setFont(font_name, font_size)

	y = page_height - margin

	def new_page() -> None:
		nonlocal y
		c.showPage()
		c.setFont(font_name, font_size)
		y = page_height - margin

	def draw_wrapped_line(line: str) -> None:
		nonlocal y
		# Wrap by words, measuring actual rendered width.
		words = line.split()
		if not words:
			y -= line_height
			if y < margin:
				new_page()
			return

		current = words[0]
		for w in words[1:]:
			test = f"{current} {w}"
			if c.stringWidth(test, font_name, font_size) <= max_width:
				current = test
			else:
				c.drawString(margin, y, current)
				y -= line_height
				if y < margin:
					new_page()
				current = w

		c.drawString(margin, y, current)
		y -= line_height
		if y < margin:
			new_page()

	for para in text.splitlines():
		para = para.rstrip()
		if not para:
			# Blank line between paragraphs.
			y -= line_height
			if y < margin:
				new_page()
			continue
		draw_wrapped_line(para)

	c.save()


def extract_text_from_docx(path: Path) -> list[str]:
	try:
		from docx import Document
	except Exception as exc:  # pragma: no cover
		raise RuntimeError("python-docx not installed. Install translate/requirements.txt.") from exc

	doc = Document(str(path))
	return [p.text for p in doc.paragraphs]


def write_docx(path: Path, paragraphs: list[str]) -> None:
	from docx import Document

	doc = Document()
	for p in paragraphs:
		doc.add_paragraph(p)
	doc.save(str(path))


def convert_doc_to_docx(doc_path: Path, *, temp_dir: Path) -> Path:
	soffice = shutil.which("soffice")
	if not soffice:
		raise RuntimeError(
			"LibreOffice CLI (soffice) not found on PATH; cannot convert .doc files."
		)

	# Convert in-place to temp_dir.
	result = subprocess.run(
		[
			soffice,
			"--headless",
			"--nologo",
			"--nolockcheck",
			"--nodefault",
			"--norestore",
			"--convert-to",
			"docx",
			"--outdir",
			str(temp_dir),
			str(doc_path),
		],
		check=False,
		capture_output=True,
		text=True,
	)
	if result.returncode != 0:
		raise RuntimeError(
			"LibreOffice conversion failed.\n"
			f"stdout:\n{result.stdout}\n\nstderr:\n{result.stderr}"
		)

	converted = temp_dir / f"{doc_path.stem}.docx"
	if not converted.exists():
		raise RuntimeError("LibreOffice reported success but output .docx was not created.")
	return converted


def out_path_for_input(
	input_path: Path,
	*,
	output_dir: Path | None,
	pdf_output: str,
) -> Path:
	base_dir = output_dir if output_dir is not None else input_path.parent

	ext = input_path.suffix.lower()
	if ext == ".pdf":
		if pdf_output == "txt":
			return base_dir / f"{input_path.stem}-eng.txt"
		if pdf_output == "docx":
			return base_dir / f"{input_path.stem}-eng.docx"
		return base_dir / f"{input_path.stem}-eng.pdf"
	if ext in {".docx", ".doc"}:
		return base_dir / f"{input_path.stem}-eng.docx"
	raise ValueError(f"Unsupported extension: {input_path.suffix}")


def translate_paragraphs(
	translator: Translator,
	paragraphs: list[str],
	*,
	source_lang: str,
	target_lang: str,
	max_chunk_chars: int,
	chunk_workers: int,
) -> list[str]:
	# Preserve empty lines by translating only non-empty paragraphs.
	out: list[str] = []
	buffer: list[str] = []

	def flush_buffer() -> None:
		if not buffer:
			return
		joined = "\n".join(buffer)
		chunks = chunk_text(joined, max_chunk_chars)
		workers = max(1, chunk_workers)
		with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
			translated_chunks = list(
				executor.map(
					lambda c: translator.translate(
						text=c, source_lang=source_lang, target_lang=target_lang
					),
					chunks,
				)
			)
		translated = "\n".join(translated_chunks).splitlines()

		# Best-effort: map translated lines back to paragraph count; if mismatch,
		# collapse to a single paragraph to avoid misalignment.
		if len(translated) != len(buffer):
			out.extend(["\n".join(translated).strip()])
		else:
			out.extend(translated)

		buffer.clear()

	for idx, p in enumerate(paragraphs):
		text = p.strip()
		if not text:
			flush_buffer()
			out.append("")
			continue

		buffer.append(text)

		# Flush before buffer gets too big.
		if sum(len(x) + 1 for x in buffer) > max_chunk_chars:
			flush_buffer()

	flush_buffer()
	return out


def main() -> int:
	parser = argparse.ArgumentParser(
		description="Translate .pdf/.doc/.docx files (default: Slovenian -> English) and write -eng outputs."
	)
	parser.add_argument(
		"--input-dir",
		required=True,
		help="Directory to scan recursively for files to translate.",
	)
	parser.add_argument(
		"--output-dir",
		default="",
		help="Optional output directory (default: next to each input file).",
	)
	parser.add_argument(
		"--provider",
		choices=["openai", "argos"],
		default="openai",
		help="Translation provider: openai (default) or argos (offline, requires installed model).",
	)
	parser.add_argument(
		"--openai-model",
		default="gpt-4o-mini",
		help="OpenAI model name (provider=openai).",
	)
	parser.add_argument("--source-lang", default="sl", help="Source language code (default: sl).")
	parser.add_argument("--target-lang", default="en", help="Target language code (default: en).")
	parser.add_argument(
		"--max-chunk-chars",
		type=int,
		default=3500,
		help="Max characters per translation chunk (lower avoids API limits).",
	)
	parser.add_argument(
		"--workers",
		type=int,
		default=4,
		help="Number of files to translate in parallel (default: 4).",
	)
	parser.add_argument(
		"--chunk-workers",
		type=int,
		default=4,
		help="Number of translation chunks to translate in parallel per file (default: 4).",
	)
	parser.add_argument(
		"--pdf-output",
		choices=["pdf", "txt", "docx"],
		default="pdf",
		help="For .pdf inputs, output format (default: pdf).",
	)
	parser.add_argument(
		"--overwrite",
		action="store_true",
		help="Overwrite existing -eng outputs.",
	)

	args = parser.parse_args()

	input_dir = Path(args.input_dir).expanduser().resolve()
	if not input_dir.exists() or not input_dir.is_dir():
		eprint(f"Input dir not found or not a directory: {input_dir}")
		return 2

	output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else None
	if output_dir is not None:
		output_dir.mkdir(parents=True, exist_ok=True)

	files = list(iter_files(input_dir))
	if not files:
		print("No supported files found.")
		return 0

	workers = max(1, args.workers)
	chunk_workers = max(1, args.chunk_workers)
	print(f"Found {len(files)} files to translate. workers={workers} chunk_workers={chunk_workers}")

	def translate_one_file(path: Path) -> str:
		translator = get_translator(args.provider, args.openai_model)
		out_path = out_path_for_input(path, output_dir=output_dir, pdf_output=args.pdf_output)
		if out_path.exists() and not args.overwrite:
			return f"Skip (exists): {out_path}"

		ext = path.suffix.lower()
		if ext == ".pdf":
			text = extract_text_from_pdf(path)
			if not text.strip():
				return f"Skip (no extractable text): {path}"
			chunks = chunk_text(text, args.max_chunk_chars)
			with concurrent.futures.ThreadPoolExecutor(max_workers=chunk_workers) as executor:
				translated_chunks = list(
					executor.map(
						lambda c: translator.translate(
							text=c, source_lang=args.source_lang, target_lang=args.target_lang
						),
						chunks,
					)
				)
			translated_text = "\n\n".join(translated_chunks).strip()
			if args.pdf_output == "txt":
				out_path.write_text(translated_text + "\n", encoding="utf-8")
			elif args.pdf_output == "docx":
				write_docx(out_path, translated_text.splitlines())
			else:
				write_pdf_from_text(out_path, translated_text)
			return f"Wrote: {out_path}"

		if ext == ".docx":
			paragraphs = extract_text_from_docx(path)
			translated = translate_paragraphs(
				translator,
				paragraphs,
				source_lang=args.source_lang,
				target_lang=args.target_lang,
				max_chunk_chars=args.max_chunk_chars,
				chunk_workers=chunk_workers,
			)
			write_docx(out_path, translated)
			return f"Wrote: {out_path}"

		if ext == ".doc":
			with tempfile.TemporaryDirectory(prefix="translate-doc-") as td:
				temp_dir = Path(td)
				docx_path = convert_doc_to_docx(path, temp_dir=temp_dir)
				paragraphs = extract_text_from_docx(docx_path)
				translated = translate_paragraphs(
					translator,
					paragraphs,
					source_lang=args.source_lang,
					target_lang=args.target_lang,
					max_chunk_chars=args.max_chunk_chars,
						chunk_workers=chunk_workers,
				)
				write_docx(out_path, translated)
			return f"Wrote: {out_path}"

		return f"Skip (unsupported): {path}"

	with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
		future_to_path = {executor.submit(translate_one_file, p): p for p in files}
		for future in concurrent.futures.as_completed(future_to_path):
			path = future_to_path[future]
			try:
				print(f"{path}: {future.result()}")
			except Exception as exc:
				eprint(f"Error translating {path}: {exc}")

	return 0


if __name__ == "__main__":
	raise SystemExit(main())


